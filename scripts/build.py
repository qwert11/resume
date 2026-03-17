from pathlib import Path
import argparse
import shutil
from jinja2 import Environment, FileSystemLoader
from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas

from utils import load_yaml, filter_items, has_any_tag, ensure_dir

ROOT = Path(__file__).resolve().parents[1]
DATA = ROOT / "data"
TARGETS = ROOT / "targets"
TEMPLATES = ROOT / "templates"
ASSETS = ROOT / "assets"
OUTPUT = ROOT / "output"

def render_markdown(context, out_path):
    env = Environment(loader=FileSystemLoader(str(TEMPLATES)))
    tpl = env.get_template("resume.md.j2")
    out_path.write_text(tpl.render(**context), encoding="utf-8")

def render_docx(context, out_path):
    doc = Document()
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(10.5)

    doc.add_heading(context["profile"]["name"], 0)
    p = doc.add_paragraph()
    p.add_run(context["target"]["title"]).bold = True
    doc.add_paragraph(context["summary"])

    doc.add_heading("Key skills", level=1)
    for skill in context["highlight_skills"]:
        doc.add_paragraph(skill, style="List Bullet")

    doc.add_heading("Achievements", level=1)
    for item in context["achievements"]:
        doc.add_paragraph(item["text"], style="List Bullet")

    doc.add_heading("Projects", level=1)
    for pitem in context["projects"]:
        p = doc.add_paragraph()
        p.add_run(f'{pitem["name"]} ({pitem["period"]})').bold = True
        doc.add_paragraph(f'Role: {pitem["role"]}')
        doc.add_paragraph(pitem["description"])

    doc.add_heading("Experience", level=1)
    for item in context["experiences"]:
        p = doc.add_paragraph()
        p.add_run(f'{item["company"]} | {item["title"]}').bold = True
        doc.add_paragraph(f'{item["start"]} - {item["end"]}')
        for bullet in item.get("bullets", []):
            doc.add_paragraph(bullet, style="List Bullet")

    doc.add_heading("Education", level=1)
    for e in context["education"]:
        doc.add_paragraph(f'{e["specialty"]} — {e["institution"]} ({e["period"]})', style="List Bullet")

    doc.save(out_path)

def render_pdf(context, out_path):
    c = canvas.Canvas(str(out_path), pagesize=A4)
    width, height = A4
    x = 40
    y = height - 50

    def line(text, size=10, leading=15, bold=False):
        nonlocal y
        font = "Helvetica-Bold" if bold else "Helvetica"
        c.setFont(font, size)
        for chunk in wrap_text(text, 105):
            c.drawString(x, y, chunk)
            y -= leading
            if y < 50:
                c.showPage()
                y = height - 50
                c.setFont(font, size)

    line(context["profile"]["name"], size=18, leading=24, bold=True)
    line(context["target"]["title"], size=12, leading=18, bold=True)
    line(context["summary"], size=10, leading=14)
    y -= 6

    line("Key skills", size=13, leading=18, bold=True)
    for s in context["highlight_skills"]:
        line(f"- {s}", size=10)

    y -= 6
    line("Achievements", size=13, leading=18, bold=True)
    for item in context["achievements"]:
        line(f"- {item['text']}", size=10)

    y -= 6
    line("Projects", size=13, leading=18, bold=True)
    for p in context["projects"]:
        line(f"{p['name']} ({p['period']})", size=11, leading=16, bold=True)
        line(f"Role: {p['role']}", size=10)
        line(p["description"], size=10)

    y -= 6
    line("Experience", size=13, leading=18, bold=True)
    for item in context["experiences"]:
        line(f"{item['company']} | {item['title']}", size=11, leading=16, bold=True)
        line(f"{item['start']} - {item['end']}", size=10)
        for bullet in item.get("bullets", []):
            line(f"• {bullet}", size=10)

    y -= 6
    line("Education", size=13, leading=18, bold=True)
    for e in context["education"]:
        line(f"- {e['specialty']} — {e['institution']} ({e['period']})", size=10)

    c.save()

def wrap_text(text, width):
    if len(text) <= width:
        return [text]
    words = text.split()
    lines = []
    current = []
    current_len = 0
    for w in words:
        add_len = len(w) + (1 if current else 0)
        if current_len + add_len <= width:
            current.append(w)
            current_len += add_len
        else:
            lines.append(" ".join(current))
            current = [w]
            current_len = len(w)
    if current:
        lines.append(" ".join(current))
    return lines

def build_target(target_id):
    profile = load_yaml(DATA / "master.yaml")["profile"]
    experiences = load_yaml(DATA / "experience.yaml")
    projects = load_yaml(DATA / "projects.yaml")
    skills = load_yaml(DATA / "skills.yaml")
    education = load_yaml(DATA / "education.yaml")
    achievements = load_yaml(DATA / "achievements.yaml")
    target = load_yaml(TARGETS / f"{target_id}.yaml")

    tags = target["include_tags"]
    exp_f = filter_items(experiences, tags)
    proj_f = filter_items(projects, tags)
    ach_f = filter_items(achievements, tags)

    # collect highlight skills
    highlight_skills = []
    for group_items in skills.values():
        for item in group_items:
            if has_any_tag(item.get("tags", []), tags):
                name = item["name"]
                if name not in highlight_skills:
                    highlight_skills.append(name)

    out_dir = OUTPUT / target_id
    ensure_dir(out_dir)

    # routing helpers for nested pages
    root_prefix = "../" if target_id != "full" else ""
    assets_prefix = "../" if target_id != "full" else ""
    downloads_prefix = "./" if target_id == "full" else "./"
    route_base = "/resume"
    page_title = f"{profile['name']} | {target['title']}"

    context = {
        "profile": profile,
        "target": target,
        "summary": profile["summary"][target["summary_key"]],
        "experiences": exp_f,
        "projects": proj_f,
        "education": education,
        "achievements": ach_f,
        "highlight_skills": highlight_skills[:14],
        "page_title": page_title,
        "root_prefix": root_prefix,
        "assets_prefix": assets_prefix,
        "downloads_prefix": downloads_prefix,
        "route_base": route_base,
    }

    # html
    env = Environment(loader=FileSystemLoader(str(TEMPLATES)))
    tpl = env.get_template("site.html.j2")
    (out_dir / "index.html").write_text(tpl.render(**context), encoding="utf-8")

    # md/docx/pdf
    render_markdown(context, out_dir / "resume.md")
    render_docx(context, out_dir / "resume.docx")
    render_pdf(context, out_dir / "resume.pdf")

def prepare_site_root():
    site_root = OUTPUT / "site"
    if site_root.exists():
        shutil.rmtree(site_root)
    ensure_dir(site_root)

    # copy assets
    shutil.copytree(ASSETS, site_root / "assets", dirs_exist_ok=True)

    # full => root index.html
    full_dir = OUTPUT / "full"
    shutil.copytree(full_dir, site_root, dirs_exist_ok=True)

    # nested routes
    for name in ["dotnet", "delphi", "web"]:
        shutil.copytree(OUTPUT / name, site_root / name, dirs_exist_ok=True)

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--target", choices=["full", "dotnet", "delphi", "web"])
    parser.add_argument("--all", action="store_true")
    args = parser.parse_args()

    if args.all:
        for t in ["full", "dotnet", "delphi", "web"]:
            build_target(t)
        prepare_site_root()
    elif args.target:
        build_target(args.target)
    else:
        raise SystemExit("Use --target or --all")

if __name__ == "__main__":
    main()
